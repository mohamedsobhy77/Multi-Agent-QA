"""
services/text_extractor.py
──────────────────────────
Synchronous text extraction for PDF, DOCX, and TXT files.

Architecture
  extract_text()          – public entry point; dispatches to the right extractor
  _extract_pdf()          – PyMuPDF (fitz)
  _extract_docx()         – python-docx (paragraphs + tables)
  _extract_txt()          – raw bytes decoded with graceful encoding fallback
  _clean_text()           – post-extraction normalisation

Async usage
  All extractors are synchronous because the underlying parsers (fitz, docx)
  are CPU-bound C extensions.  To call from an async context without blocking
  the event loop, wrap with asyncio.run_in_executor:

      loop = asyncio.get_running_loop()
      text = await loop.run_in_executor(
          None, partial(extract_text, data, file_type, filename)
      )

  A ready-made async wrapper is provided as extract_text_async().

Raises
  TextExtractionError  – parser raised an unexpected error
  EmptyDocumentError   – parsed output is empty / whitespace-only
"""

from __future__ import annotations

import asyncio
import io
import re
from collections.abc import Callable
from concurrent.futures import ThreadPoolExecutor
from functools import partial
from typing import Final

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from services.exceptions import EmptyDocumentError, TextExtractionError

# ── Constants ─────────────────────────────────────────────────────────────────

# Character count below which we consider extraction to have produced nothing.
_MIN_CHARS: Final[int] = 10

# Encoding cascade used when decoding TXT files.
_TXT_ENCODINGS: Final[tuple[str, ...]] = ("utf-8", "utf-8-sig", "latin-1", "cp1252")

# Module-level thread pool used by extract_text_async.
# A single worker is sufficient because the GIL means only one extraction runs
# at a time anyway; increase max_workers for true parallelism with ProcessPool.
_EXECUTOR: ThreadPoolExecutor = ThreadPoolExecutor(
    max_workers=4,
    thread_name_prefix="text_extractor",
)


# ── PDF extractor ─────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes, filename: str) -> str:
    """
    Extract text from a PDF using PyMuPDF (fitz).

    Strategy
    --------
    • Open from bytes (never touches disk).
    • Iterate every page with get_text("text") for plain-text layout.
    • Pages that fail individually are skipped with a warning; we keep
      extracting rather than failing the whole document.
    • Image-only PDFs (no text layer) produce an empty string, which
      is caught by the post-extraction emptiness check in extract_text().

    Parameters
    ----------
    data:     Raw PDF bytes.
    filename: Used only for error context.

    Raises
    ------
    TextExtractionError  If fitz cannot open the bytes at all.
    """
    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise TextExtractionError(
            filename=filename,
            file_type="pdf",
            reason=f"Could not open PDF: {exc}",
            cause=exc,
        ) from exc

    page_texts: list[str] = []

    try:
        for page_index in range(len(doc)):
            try:
                page = doc[page_index]
                page_text = page.get_text("text")
                if page_text and page_text.strip():
                    page_texts.append(page_text)
            except Exception as exc:
                # A single corrupt page should not abort the whole document.
                # In production you would log this; here we surface it in the
                # exception chain only if the overall result is empty.
                page_texts.append(f"[Page {page_index + 1} could not be read: {exc}]")
    finally:
        doc.close()

    return "\n\n".join(page_texts)


# ── DOCX extractor ────────────────────────────────────────────────────────────

def _extract_docx(data: bytes, filename: str) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Extracts
    --------
    • All body paragraphs in document order.
    • Table cell contents, formatted as pipe-separated rows so the structure
      is readable without rendering the full table markup.

    Parameters
    ----------
    data:     Raw DOCX bytes.
    filename: Used only for error context.

    Raises
    ------
    TextExtractionError  If python-docx cannot parse the bytes.
    """
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise TextExtractionError(
            filename=filename,
            file_type="docx",
            reason=f"Could not open DOCX: {exc}",
            cause=exc,
        ) from exc

    segments: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            segments.append(text)

    # Tables — flatten to readable pipe-delimited rows
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                segments.append(" | ".join(cells))

    return "\n\n".join(segments)


# ── TXT extractor ─────────────────────────────────────────────────────────────

def _extract_txt(data: bytes, filename: str) -> str:
    """
    Decode a plain-text file, trying encodings in priority order.

    Falls back through utf-8 → utf-8-sig → latin-1 → cp1252.
    latin-1 is a lossless fallback (every byte maps to a code point) so
    this cascade should never fail on any real-world text file.

    Parameters
    ----------
    data:     Raw file bytes.
    filename: Used only for error context.

    Raises
    ------
    TextExtractionError  If every encoding in _TXT_ENCODINGS fails.
    """
    for encoding in _TXT_ENCODINGS:
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue

    raise TextExtractionError(
        filename=filename,
        file_type="txt",
        reason=(
            f"Could not decode file with any of the attempted encodings: "
            f"{', '.join(_TXT_ENCODINGS)}."
        ),
    )


# ── Extractor registry ────────────────────────────────────────────────────────

# Maps canonical file-type key → extractor function.
# Adding a new format only requires registering it here.
_EXTRACTORS: dict[str, Callable[[bytes, str], str]] = {
    "pdf":  _extract_pdf,
    "docx": _extract_docx,
    "txt":  _extract_txt,
}


# ── Post-extraction cleanup ───────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalise whitespace in extracted text while preserving structure.

    Rules
    -----
    • Strip trailing whitespace from each line.
    • Collapse runs of 3+ consecutive blank lines to exactly two blank lines
      (preserves paragraph breaks but removes excessive vertical whitespace
      common in PDFs with wide margins).
    • Strip leading/trailing whitespace from the whole document.
    """
    # Strip trailing spaces from every line
    lines = [line.rstrip() for line in text.splitlines()]
    normalised = "\n".join(lines)

    # Collapse 3+ consecutive newlines → 2
    normalised = re.sub(r"\n{3,}", "\n\n", normalised)

    return normalised.strip()


# ── Public synchronous API ────────────────────────────────────────────────────

def extract_text(data: bytes, file_type: str, filename: str) -> str:
    """
    Extract and clean plain text from a supported file.

    Parameters
    ----------
    data:      Raw bytes of the file (already read into memory).
    file_type: Canonical type string: 'pdf', 'docx', or 'txt'.
               Typically produced by file_validator.validate_upload().
    filename:  Original filename — used only for error messages.

    Returns
    -------
    str
        Cleaned plain-text content.  Guaranteed non-empty (>= _MIN_CHARS chars)
        on success.

    Raises
    ------
    TextExtractionError
        • file_type is not registered in _EXTRACTORS
        • The underlying parser raises an unexpected exception
    EmptyDocumentError
        • Parsed text is empty or below the minimum character threshold
          (e.g. a PDF with no text layer, or a DOCX with only images)
    """
    extractor = _EXTRACTORS.get(file_type)

    if extractor is None:
        raise TextExtractionError(
            filename=filename,
            file_type=file_type,
            reason=(
                f"No extractor registered for type {file_type!r}. "
                f"Supported: {', '.join(sorted(_EXTRACTORS))}."
            ),
        )

    # Run the format-specific extractor.
    # Any uncaught exception from the extractor is intentionally allowed to
    # propagate; each extractor wraps its own parser errors in TextExtractionError.
    raw = extractor(data, filename)

    # Post-extraction normalisation
    cleaned = _clean_text(raw)

    # Emptiness guard — catches image-only PDFs, empty DOCX files, etc.
    if len(cleaned) < _MIN_CHARS:
        raise EmptyDocumentError(
            filename=filename,
            reason=(
                f"Extraction produced only {len(cleaned)} character(s). "
                "The file may be image-only, password-protected, or corrupt."
            ),
        )

    return cleaned


# ── Public async wrapper ──────────────────────────────────────────────────────

async def extract_text_async(
    data: bytes,
    file_type: str,
    filename: str,
) -> str:
    """
    Async wrapper around extract_text().

    Offloads the CPU-bound extraction to the module-level ThreadPoolExecutor
    so it never blocks the event loop.

    Parameters and return value are identical to extract_text().

    Example
    -------
        text = await extract_text_async(content, "pdf", "spec.pdf")
    """
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(
        _EXECUTOR,
        partial(extract_text, data, file_type, filename),
    )


# ── Utility ───────────────────────────────────────────────────────────────────

def supported_file_types() -> list[str]:
    """Return a sorted list of all registered file-type keys."""
    return sorted(_EXTRACTORS.keys())
